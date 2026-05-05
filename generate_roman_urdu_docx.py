from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

def create_comprehensive_roman_urdu_doc():
    doc = Document()
    
    # Global Font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Cfab Backend - Complete API Documentation (Roman Urdu)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Yeh document backend ki saari APIs ko cover karta hai jo front-end integration ke liye zaroori hain.')
    
    doc.add_paragraph('Base URL: http://localhost:8000')

    # --- Section 1: Authentication ---
    doc.add_heading('1. Authentication (User Login & Register)', level=1)
    
    doc.add_heading('Register Naya User', level=2)
    doc.add_paragraph('Endpoint: POST /auth/register')
    doc.add_paragraph('Purpose: Naya account banane ke liye.')
    add_code_block(doc, 'Body: {\n  "email": "...",\n  "name": "...",\n  "password": "...",\n  "role": "student"\n}')

    doc.add_heading('Login User', level=2)
    doc.add_paragraph('Endpoint: POST /auth/login')
    doc.add_paragraph('Purpose: JWT Token hasil karne ke liye.')
    doc.add_paragraph('Body (Form-Data): username, password')

    # --- Section 2: Dashboard ---
    doc.add_heading('2. Dashboard (Check Yourself)', level=1)
    
    doc.add_heading('Unified Dashboard Data', level=2)
    doc.add_paragraph('Endpoint: GET /check-yourself')
    doc.add_paragraph('Purpose: Ek saath saari assignments aur quizzes ki list fetch karna.')

    # --- Section 3: Assignments ---
    doc.add_heading('3. Assignments Management', level=1)
    
    doc.add_heading('Assignments ki List', level=2)
    doc.add_paragraph('Endpoint: GET /assignments')
    
    doc.add_heading('Assignment Search', level=2)
    doc.add_paragraph('Endpoint: GET /assignments/search?title=Keyword')

    doc.add_heading('Assignment Detail', level=2)
    doc.add_paragraph('Endpoint: GET /assignments/{id}')
    doc.add_paragraph('Purpose: Assignment ki poori detail (PDF link ya Coding metadata).')

    doc.add_heading('PDF File Download', level=2)
    doc.add_paragraph('Endpoint: GET /assignments/{id}/file')

    # --- Section 4: Quizzes ---
    doc.add_heading('4. Quizzes', level=1)
    
    doc.add_heading('Quizzes ki List', level=2)
    doc.add_paragraph('Endpoint: GET /quiz')

    doc.add_heading('Quiz Submit Karein', level=2)
    doc.add_paragraph('Endpoint: POST /quiz/submit')
    add_code_block(doc, 'Body: {\n  "quiz_id": "...",\n  "answers": [\n    {"question_id": "...", "selected": "A"}\n  ]\n}')

    # --- Section 5: Coding Practice ---
    doc.add_heading('5. Coding Practice (Evaluation)', level=1)
    
    doc.add_heading('Code Submit Karein', level=2)
    doc.add_paragraph('Endpoint: POST /practice-code/submit')
    add_code_block(doc, 'Body: {\n  "assignment_id": "...",\n  "code": "...",\n  "language": "python"\n}')

    doc.add_heading('Result Status Check', level=2)
    doc.add_paragraph('Endpoint: GET /practice-code/{submission_id}')
    doc.add_paragraph('Note: Is endpoint ko poll karna hota hai (PENDING -> SUCCESS/ERROR).')

    doc.add_heading('My Submissions', level=2)
    doc.add_paragraph('Endpoint: GET /practice-code/me/list')

    # --- Section 6: Admin Actions ---
    doc.add_heading('6. Admin Panels', level=1)
    
    doc.add_heading('Upload PDF Assignment', level=2)
    doc.add_paragraph('Endpoint: POST /assignments/upload (Form-Data)')

    doc.add_heading('Create Coding Assignment', level=2)
    doc.add_paragraph('Endpoint: POST /assignments/coding (JSON)')

    doc.add_heading('Quiz PDF Parsing', level=2)
    doc.add_paragraph('Endpoint: POST /admin/quiz/upload-pdf (Form-Data)')

    doc.add_heading('Delete Endpoints', level=2)
    doc.add_paragraph('- DELETE /assignments/{id}\n- DELETE /admin/quiz/{id}\n- DELETE /admin/coding-assignment/{id}')

    # Save
    filename = 'API_DOCUMENTATION_ROMAN_URDU.docx'
    doc.save(filename)
    print(f"Successfully updated '{filename}'!")

if __name__ == "__main__":
    create_comprehensive_roman_urdu_doc()
