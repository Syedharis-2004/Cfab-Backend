import asyncio
from docx import Document
from docx.shared import Pt, RGBColor
from app.main import app

def create_comprehensive_api_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Cfab Backend - Comprehensive API Documentation', 0)
    
    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph('Base URL: http://localhost:8000')
    doc.add_paragraph('Interactive Swagger Docs: http://localhost:8000/docs')

    routes_by_tag = {}

    for route in app.routes:
        # We only care about APIRoutes
        if hasattr(route, 'methods'):
            tags = getattr(route, 'tags', [])
            tag = tags[0] if tags else "General"
            
            if tag not in routes_by_tag:
                routes_by_tag[tag] = []
            
            methods = ", ".join(route.methods - {"OPTIONS", "HEAD"}) if route.methods else "GET"
            path = route.path
            name = route.name
            description = getattr(route, 'description', '') or route.endpoint.__doc__ or 'No description provided.'
            description = description.strip()
            
            routes_by_tag[tag].append({
                "methods": methods,
                "path": path,
                "name": name,
                "description": description
            })

    for tag, routes in routes_by_tag.items():
        doc.add_heading(f'{tag.capitalize()} Endpoints', level=1)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Method'
        hdr_cells[1].text = 'Endpoint'
        hdr_cells[2].text = 'Description / Purpose'
        
        for r in hdr_cells:
            for paragraph in r.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        for route in routes:
            row_cells = table.add_row().cells
            row_cells[0].text = route['methods']
            row_cells[1].text = route['path']
            row_cells[2].text = route['description']

    filename = 'Comprehensive_API_Routes_v2.docx'
    doc.save(filename)
    print(f"Successfully created '{filename}'!")

if __name__ == "__main__":
    create_comprehensive_api_doc()
