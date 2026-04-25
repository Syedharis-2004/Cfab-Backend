try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Error: 'python-docx' library not found. Please run 'pip install python-docx' first.")
    exit(1)

def create_api_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Check Yourself API Documentation', 0)
    
    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph('Base URL: http://localhost:8000')
    doc.add_paragraph('Interactive Swagger Docs: http://localhost:8000/docs')

    # Auth
    doc.add_heading('🔐 Authentication', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Method'
    hdr_cells[1].text = 'Endpoint'
    hdr_cells[2].text = 'Description'
    
    auth_routes = [
        ('POST', '/auth/register', 'Register a new user'),
        ('POST', '/auth/login', 'Login and get JWT token'),
    ]
    for m, e, d in auth_routes:
        row_cells = table.add_row().cells
        row_cells[0].text = m
        row_cells[1].text = e
        row_cells[2].text = d

    # Assignments
    doc.add_heading('📚 Assignments', level=1)
    doc.add_paragraph('GET /assignments/ - List all')
    doc.add_paragraph('GET /assignments/{id} - View details')
    doc.add_paragraph('POST /assignments/{id}/submit - Submit PDF')

    # Quizzes
    doc.add_heading('📝 Quizzes', level=1)
    doc.add_paragraph('GET /quiz/ - List all')
    doc.add_paragraph('POST /quiz/{id}/submit - Submit answers')

    # Coding
    doc.add_heading('💻 Coding Practice (Sandbox)', level=1)
    doc.add_paragraph('POST /submissions/ - Submit code for evaluation')
    doc.add_paragraph('GET /submissions/{id} - Get results (Async)')

    doc.save('API_ROUTES.docx')
    print("Successfully created 'API_ROUTES.docx'!")

if __name__ == "__main__":
    create_api_doc()
