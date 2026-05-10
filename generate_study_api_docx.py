from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

def create_study_api_doc():
    doc = Document()
    
    # Global Font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Study aur Time Management API Documentation (Roman Urdu)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Yeh document aapko Study module aur Time Management module ki APIs ke baare mein batayega ke kaunsi API kya kaam karti hai aur unka request structure kya hai.')
    
    doc.add_paragraph('Base URL: http://localhost:8000')

    # --- Section 1: Study Module ---
    doc.add_heading('1. Study Module (Course aur Lecture Management)', level=1)
    
    doc.add_heading('Admin: Course Banayein', level=2)
    doc.add_paragraph('Endpoint: POST /api/study/course/create')
    doc.add_paragraph('Purpose: Naya course banane ke liye.')
    add_code_block(doc, '{\n  "title": "FastAPI Masterclass",\n  "description": "Learn FastAPI from scratch"\n}')

    doc.add_heading('Admin: Lecture Upload Karein', level=2)
    doc.add_paragraph('Endpoint: POST /api/study/lecture/upload')
    doc.add_paragraph('Purpose: Kisi course mein lecture add karne ke liye.')
    add_code_block(doc, '{\n  "course_id": "course_id_here",\n  "title": "Introduction to API",\n  "video_url": "https://youtube.com/link",\n  "duration": 45,\n  "order_index": 1\n}')

    doc.add_heading('Student/Admin: Course ki Lectures Dekhein', level=2)
    doc.add_paragraph('Endpoint: GET /api/study/lecture/course/{course_id}')
    doc.add_paragraph('Purpose: Kisi course ki saari lectures list karne ke liye.')

    # --- Section 2: Time Management Module ---
    doc.add_heading('2. Time Management Module (Student Dashboard)', level=1)
    
    doc.add_heading('Time Management Dashboard', level=2)
    doc.add_paragraph('Endpoint: GET /api/time-management')
    doc.add_paragraph('Purpose: Student dashboard jahan active aur available courses dikhte hain.')

    doc.add_heading('Course Shuru Karein (Generate Plan)', level=2)
    doc.add_paragraph('Endpoint: POST /api/time-management/course/start')
    doc.add_paragraph('Purpose: Student ki availability ke hisaab se schedule banane ke liye.')
    add_code_block(doc, '{\n  "course_id": "course_id_here",\n  "daily_minutes": 60,\n  "weekly_schedule": {\n    "monday": 60,\n    "tuesday": 30,\n    "wednesday": 90\n  }\n}')

    doc.add_heading('Study Plan Dekhein', level=2)
    doc.add_paragraph('Endpoint: GET /api/time-management/study-plan/{course_id}')
    doc.add_paragraph('Purpose: Student ka locked schedule dekhne ke liye.')

    doc.add_heading('Lecture Complete Karein', level=2)
    doc.add_paragraph('Endpoint: POST /api/time-management/lecture/complete')
    doc.add_paragraph('Purpose: Lecture khatam hone par progress update karne ke liye.')
    add_code_block(doc, '{\n  "course_id": "course_id_here",\n  "lecture_id": "lecture_id_here"\n}')

    doc.add_heading('Progress Check Karein', level=2)
    doc.add_paragraph('Endpoint: GET /api/time-management/progress/{course_id}')

    # --- Section 3: Aham Points ---
    doc.add_heading('Aham Points (Frontend Tips)', level=1)
    doc.add_paragraph('1. ID Handling: Hamesha id field use karein, jo ab string format mein hai.')
    doc.add_paragraph('2. Hydrated Data: Study Plan API ab lecture ki IDs nahi balki poore objects deti hai.')
    doc.add_paragraph('3. Authentication: Header mein Authorization: Bearer <token> bhejna zaroori hai.')

    # Save
    filename = 'API_DOCUMENTATION_STUDY_TIME_MANAGEMENT.docx'
    doc.save(filename)
    print(f"Successfully created '{filename}'!")

if __name__ == "__main__":
    create_study_api_doc()
